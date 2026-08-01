"""OCR 文本/视觉表格 → 可编辑 Word 表格（供 DOCX 组装脚本导入）。

功能:
- parse_ascii_table: 把 OCR 输出的 | 分隔或 ┌┬┐ 框线表格解析为二维数组
- add_table_to_doc: 把二维数组写入 python-docx 文档（Table Grid，列宽自适应）
- add_json_tables: 把 table_extract.py 输出的 JSON 表格写入文档

命令行自测:
    .venv\\Scripts\\python.exe scripts\\docx_tables.py --input page.txt --output 表格样例.docx
"""

from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

BODY_CN = "宋体"
BODY_EN = "Times New Roman"
TABLE_STYLE = "Table Grid"
TABLE_WIDTH_CM = 16.0


def _is_border_line(cells: list[str]) -> bool:
    return all(not re.search(r"[\u4e00-\u9fffA-Za-z0-9ɕʑɿʅŋɑɛɔʊɤæəɪʊ]", c or "") for c in cells)


def parse_ascii_table(lines: list[str]) -> list[list[str]]:
    """把连续的表格式文本行解析为二维数组（支持 | 分隔与框线表格）。"""
    rows: list[list[str]] = []
    for raw in lines:
        line = raw.strip()
        if not line or ("|" not in line and "│" not in line):
            continue
        if "│" in line:
            inner = line.strip("│").strip()
            cells = [c.strip() for c in inner.split("│")]
        else:
            inner = line.strip("|").strip()
            cells = [c.strip() for c in inner.split("|")]
        if _is_border_line(cells):
            continue
        if any(cells):
            rows.append(cells)
    return rows


def _set_cell_run(run, size: float, bold: bool = False, color=None):
    run.font.name = BODY_EN
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_CN)
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_table_to_doc(
    doc: Document,
    rows: list[list[str]],
    *,
    font_size: float = 9.5,
    header_bold: bool = True,
    max_width_cm: float = TABLE_WIDTH_CM,
) -> object:
    """把二维数组写入 Word 文档，返回 python-docx Table 对象。"""
    if not rows:
        return None
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = TABLE_STYLE
    table.autofit = False
    table.alignment = 1  # CENTER

    # 列宽按内容长度比例分配
    col_max = [0] * ncols
    for row in rows:
        for j, cell_text in enumerate(row):
            col_max[j] = max(col_max[j], len(cell_text or ""))
    total = max(sum(col_max), 1)
    widths = [max(max_width_cm * (m / total), 1.2) for m in col_max]
    width_sum = sum(widths)
    if width_sum > max_width_cm:
        widths = [w * max_width_cm / width_sum for w in widths]

    for i, row in enumerate(rows):
        for j in range(ncols):
            cell = table.cell(i, j)
            cell.width = Cm(widths[j])
            text = row[j] if j < len(row) else ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            for part in re.split(r"(\[[?？]\])", text):
                if not part:
                    continue
                run = p.add_run(part)
                if re.fullmatch(r"\[[?？]\]", part):
                    _set_cell_run(run, font_size, bold=True, color=(255, 0, 0))
                else:
                    _set_cell_run(run, font_size, bold=(i == 0 and header_bold))
    return table


def add_json_tables(doc: Document, tables: list[dict]) -> None:
    """把 table_extract.py 输出的 JSON 表格写入文档。"""
    for table in tables:
        rows = table.get("rows", [])
        add_table_to_doc(doc, rows)
        doc.add_paragraph().paragraph_format.space_after = Pt(4)


def group_table_lines(lines: list[str]) -> list[list[str]]:
    """把文本行中连续的表格行分组，返回每组行列表。"""
    groups: list[list[str]] = []
    current: list[str] = []
    for line in lines:
        if "|" in line or "│" in line:
            current.append(line)
        else:
            if current:
                groups.append(current)
                current = []
    if current:
        groups.append(current)
    return groups


def main() -> int:
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    parser = argparse.ArgumentParser(description="OCR 表格 → Word 表格（自测/独立转换）")
    parser.add_argument("--input", required=True, help="包含表格的文本文件")
    parser.add_argument("--output", required=True, help="输出 DOCX")
    parser.add_argument("--tables-json", default="", help="可选：table_extract 输出的 JSON")
    args = parser.parse_args()

    doc = Document()
    lines = Path(args.input).read_text(encoding="utf-8").splitlines()
    converted = 0
    for group in group_table_lines(lines):
        rows = parse_ascii_table(group)
        if rows:
            add_table_to_doc(doc, rows)
            converted += 1
            doc.add_paragraph().paragraph_format.space_after = Pt(4)
    if args.tables_json:
        data = json.loads(Path(args.tables_json).read_text(encoding="utf-8"))
        for page_data in data.values():
            add_json_tables(doc, page_data.get("tables", []))
            converted += len(page_data.get("tables", []))
    Path(args.output).parent.mkdir(parents=True, exist_ok=True)
    doc.save(args.output)
    print(f"已转换 {converted} 个表格 -> {args.output}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
